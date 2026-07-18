"""DOCX -> plain text (with hyperlink URLs preserved).

A ``.docx`` is a zip of XML, and python-docx hands us the paragraphs directly.
Two things a naive ``document.paragraphs`` dump gets wrong, both of which matter
for CVs, and both handled here:

  * **Tables.** Table text is *not* in ``document.paragraphs`` - it lives in a
    parallel ``document.tables`` list. CV templates lean on tables for layout
    (a left "dates" column beside a right "description" column), so we walk the
    document body in order and pull text from paragraphs and table cells alike
    (including nested tables).

  * **Hyperlink URLs.** ``paragraph.text`` yields only the visible anchor text,
    never the underlying href - so a text-only dump turns "GitHub" into dead
    text and loses the URL the CV must keep clickable. We reconstruct each
    paragraph from ``iter_inner_content()`` (runs and hyperlinks in order) and
    splice the URL in next to its anchor as ``anchor <url>``, so the LLM step
    sees which link belongs to which label. :func:`extract_hyperlinks` returns
    the same links as structured ``(text, url)`` pairs.

Still not captured (documented, not hidden): headers, footers, and text boxes -
they live outside the document body.

Run the smoke test with the package form so ``docx`` still resolves to the
installed library (a direct ``python .../docx.py`` would import this file as
``docx`` and shadow it)::

    python -m cv_agent.parsers.docx [path/to/cv.docx]
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterator, List, Optional, Set, Tuple, Union

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph

# A "parent" we can walk block-by-block: the whole document, or one table cell.
_BlockParent = Union[_Document, _Cell]

_BLANK_RUN = re.compile(r"\n{3,}")


def _open(source: Union[str, Path]) -> _Document:
    """Open ``source`` as a docx, with clear errors for the common mistakes."""
    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(f"No such file: {path}")
    try:
        return Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(
            f"{path} is not a valid .docx file (is it a .doc, PDF, or corrupt?)."
        ) from exc


def _iter_block_items(parent: _BlockParent) -> Iterator[Union[Paragraph, Table]]:
    """Yield each paragraph and table under ``parent``, in document order.

    python-docx exposes paragraphs and tables as two separate lists, which loses
    their interleaving. This walks the underlying XML so a table that sits
    between two paragraphs comes out in the right place.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:  # pragma: no cover - defensive
        raise TypeError(f"Cannot iterate block items of {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _iter_paragraphs(
    parent: _BlockParent,
    include_tables: bool,
    _seen: Optional[Set[int]] = None,
) -> Iterator[Paragraph]:
    """Yield every paragraph under ``parent`` depth-first, table cells included.

    ``_seen`` dedups by element identity: a horizontally merged table cell is
    reported once per grid column by ``row.cells``, which would otherwise repeat
    its paragraphs (and their links). Skipping already-seen ``<w:p>`` elements
    drops those repeats without dropping genuinely distinct content.
    """
    if _seen is None:
        _seen = set()
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            key = id(block._p)
            if key in _seen:
                continue
            _seen.add(key)
            yield block
        elif isinstance(block, Table) and include_tables:
            for row in block.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell, include_tables, _seen)


def _paragraph_text(paragraph: Paragraph, annotate_links: bool) -> str:
    """Text of one paragraph, optionally with ``anchor <url>`` for hyperlinks."""
    if not annotate_links:
        return paragraph.text
    parts: List[str] = []
    for item in paragraph.iter_inner_content():
        if isinstance(item, Hyperlink):
            # .url is the full external href ('' for internal-only bookmarks).
            parts.append(f"{item.text} <{item.url}>" if item.url else item.text)
        else:  # a Run (or any future inline content) - just take its text
            parts.append(getattr(item, "text", ""))
    return "".join(parts)


def extract_text(
    source: Union[str, Path],
    *,
    include_tables: bool = True,
    annotate_links: bool = True,
) -> str:
    """Extract all body text from a ``.docx`` file as a single string.

    Args:
        source: path to the ``.docx`` file.
        include_tables: also pull text from table cells (default True). Turn off
            only if you specifically want top-level paragraphs alone.
        annotate_links: splice each hyperlink's URL in next to its anchor text as
            ``anchor <url>`` (default True) so links survive the text dump. Set
            False for anchor text only.

    Returns:
        The document's text, one paragraph per line, with runs of blank lines
        collapsed to a single blank line so section breaks survive but noise
        (e.g. empty table cells) does not.

    Raises:
        FileNotFoundError: if ``source`` does not exist.
        ValueError: if ``source`` is not a readable ``.docx`` package.
    """
    document = _open(source)
    lines = [
        _paragraph_text(p, annotate_links).rstrip()
        for p in _iter_paragraphs(document, include_tables)
    ]
    text = "\n".join(lines)
    return _BLANK_RUN.sub("\n\n", text).strip()


def extract_hyperlinks(
    source: Union[str, Path],
    *,
    include_tables: bool = True,
) -> List[Tuple[str, str]]:
    """Return every external hyperlink as ``(anchor_text, url)``, in order.

    Internal-only links (bookmarks with no external target) are skipped. Handy
    for mapping links straight onto the schema's ``Link(label, url)`` later.
    """
    document = _open(source)
    links: List[Tuple[str, str]] = []
    for paragraph in _iter_paragraphs(document, include_tables):
        for hl in paragraph.hyperlinks:
            if hl.url:
                links.append((hl.text, hl.url))
    return links


if __name__ == "__main__":  # pragma: no cover
    import sys

    # Windows consoles default to cp1252 and choke on non-Latin CV text; force
    # UTF-8 so the dump prints faithfully instead of raising UnicodeEncodeError.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

    if len(sys.argv) < 2:
        print("usage: python -m cv_agent.parsers.docx <path/to/cv.docx>")
        raise SystemExit(2)

    path = sys.argv[1]
    text = extract_text(path)
    links = extract_hyperlinks(path)

    print(f"--- extracted {len(text)} chars, {len(links)} hyperlink(s) ---\n")
    print(text)
    if links:
        print("\n--- hyperlinks ---")
        for anchor, url in links:
            print(f"  {anchor!r} -> {url}")
